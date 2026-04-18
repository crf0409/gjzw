#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""将实验总结报告 Markdown 转为 .docx 格式（含图表嵌入）"""

import os
import json
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

PAPER_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                         'outputs', 'paper')
FIG_DIR = os.path.join(PAPER_DIR, 'figures')


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elm)


def add_styled_table(doc, headers, rows, col_widths=None, highlight_best_col=None):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    # 数据行
    best_row_idx = None
    if highlight_best_col is not None:
        try:
            vals = [float(r[highlight_best_col]) for r in rows]
            best_row_idx = vals.index(max(vals))
        except (ValueError, IndexError):
            pass

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)

        if best_row_idx is not None and r_idx == best_row_idx:
            for c_idx in range(len(headers)):
                set_cell_shading(table.rows[r_idx + 1].cells[c_idx], 'D5F5E3')

    return table


def add_figure(doc, filename, caption, width=Inches(5.5)):
    """添加图片及说明"""
    path = os.path.join(FIG_DIR, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.style = doc.styles['Caption'] if 'Caption' in [s.name for s in doc.styles] else None
        for run in cap.runs:
            run.font.size = Pt(9)
            run.italic = True
    else:
        doc.add_paragraph(f'[图片缺失: {filename}]')


def main():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ── 标题页 ──
    title = doc.add_heading('基于深度学习的古建筑文字图像分类', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('实验总结报告', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('实验框架：PyTorch 2.9.1 | GPU: NVIDIA RTX 3090 × 4\n').font.size = Pt(11)
    info.add_run('报告日期：2026年2月10日').font.size = Pt(11)

    doc.add_page_break()

    # ── 1. 实验概述 ──
    doc.add_heading('1. 实验概述', level=1)

    doc.add_heading('1.1 研究背景与目标', level=2)
    doc.add_paragraph(
        '本实验旨在系统性评估多种主流深度学习模型在古建筑文字图像分类任务上的性能表现。'
        '通过对比分析9种不同架构的分类模型，从分类精度、模型效率、计算成本等多个维度进行全面评估，'
        '为古建筑文字识别领域的模型选择提供科学依据。'
    )

    doc.add_heading('1.2 实验环境', level=2)
    add_styled_table(doc,
        ['项目', '配置'],
        [
            ['操作系统', 'Linux 6.14.0-37-generic'],
            ['深度学习框架', 'PyTorch 2.9.1 + CUDA 12.8'],
            ['GPU', 'NVIDIA GeForce RTX 3090 (24GB) × 4'],
            ['Python', '3.11'],
            ['辅助库', 'torchvision, timm, scikit-learn, matplotlib'],
        ]
    )

    # ── 2. 数据集描述 ──
    doc.add_heading('2. 数据集描述', level=1)

    doc.add_heading('2.1 数据规模', level=2)
    doc.add_paragraph(
        '总计2,269张古建筑文字图像，划分为训练集（1,815张）和独立测试集（454张）。'
        '训练集进一步按7:3比例划分为训练子集（1,270张）和验证子集（545张），采用分层采样保持类别分布一致。'
    )

    doc.add_heading('2.2 类别分布', level=2)
    add_styled_table(doc,
        ['类别', '训练集', '测试集', '合计', '占比'],
        [
            ['类别 1', '264', '60', '324', '14.3%'],
            ['类别 2', '502', '134', '636', '28.0%'],
            ['类别 3', '336', '82', '418', '18.4%'],
            ['类别 4', '126', '32', '158', '7.0%'],
            ['类别 5', '232', '58', '290', '12.8%'],
            ['类别 6', '355', '88', '443', '19.5%'],
        ]
    )

    doc.add_paragraph(
        '数据集存在一定类别不平衡，类别4仅占7.0%而类别2占28.0%。'
        '训练时采用加权交叉熵损失函数（Weighted Cross-Entropy Loss）缓解不平衡问题。'
    )

    doc.add_heading('2.3 数据预处理与增强', level=2)
    doc.add_paragraph('预处理流程：灰度加载 → 朝向矫正 → 尺寸调整 → 灰度转RGB（预训练模型需要）')
    add_styled_table(doc,
        ['增强方式', '参数'],
        [
            ['随机旋转', '±54°'],
            ['随机平移', '水平/垂直各 ±10%'],
            ['随机缩放', '0.85x ~ 1.15x'],
            ['亮度抖动', '±10%'],
            ['对比度抖动', '0.9 ~ 1.1'],
        ]
    )

    # ── 3. 实验模型 ──
    doc.add_heading('3. 实验模型', level=1)
    doc.add_paragraph('共评估9种模型架构，涵盖经典CNN、轻量级网络、高效网络和Transformer架构：')

    # 加载性能指标
    perf = {}
    pf = os.path.join(FIG_DIR, 'model_perf_metrics.json')
    if os.path.exists(pf):
        with open(pf) as f:
            perf = json.load(f)

    add_styled_table(doc,
        ['模型', '类型', '输入尺寸', '总参数量', '预训练'],
        [
            ['Custom MLP', '全连接网络', '224×224', '12.88M', '无'],
            ['ResNet-50', '残差网络', '224×224', '24.04M', 'ImageNet'],
            ['VGG-16', '经典CNN', '224×224', '14.85M', 'ImageNet'],
            ['VGG-19', '经典CNN', '224×224', '20.16M', 'ImageNet'],
            ['Inception V3', 'Inception系列', '299×299', '22.32M', 'ImageNet'],
            ['Inception-ResNet V2', 'Inception+残差', '299×299', '54.70M', 'ImageNet'],
            ['EfficientNet-B3', '高效网络', '300×300', '11.09M', 'ImageNet'],
            ['MobileNet V3', '轻量级网络', '224×224', '3.22M', 'ImageNet'],
            ['ViT-B/16', 'Transformer', '224×224', '86.00M', 'ImageNet'],
        ]
    )

    doc.add_heading('3.1 迁移学习策略', level=2)
    doc.add_paragraph(
        '所有预训练模型采用部分冻结微调（Partial Fine-tuning）策略：前80%参数冻结，仅微调后20%参数。'
        '原始分类器替换为统一的自定义分类头：BatchNorm → Dropout(0.3) → FC(dim, 256) → ReLU → Dropout(0.2) → FC(256, 6)。'
    )

    # ── 4. 训练配置 ──
    doc.add_heading('4. 训练配置', level=1)
    add_styled_table(doc,
        ['参数', '值'],
        [
            ['优化器', 'Adam (weight_decay=0.0001)'],
            ['初始学习率', '0.0001'],
            ['学习率调度', 'Cosine Annealing + ReduceLROnPlateau'],
            ['批次大小', '32'],
            ['最大训练轮数', '100'],
            ['早停策略', 'patience=15（基于验证集准确率）'],
            ['损失函数', '加权交叉熵（Weighted CrossEntropyLoss）'],
            ['随机种子', '42'],
        ]
    )

    doc.add_heading('4.1 训练过程概览', level=2)
    add_styled_table(doc,
        ['模型', '训练轮数', '训练时间', '早停'],
        [
            ['Custom MLP', '52', '2.2 min', '是'],
            ['ResNet-50', '41', '2.8 min', '是'],
            ['VGG-16', '41', '3.5 min', '是'],
            ['VGG-19', '59', '6.5 min', '是'],
            ['Inception V3', '46', '4.0 min', '是'],
            ['Inception-ResNet V2', '25', '3.5 min', '是'],
            ['EfficientNet-B3', '42', '3.7 min', '是'],
            ['MobileNet V3', '38', '2.2 min', '是'],
            ['ViT-B/16', '59', '8.6 min', '是'],
        ]
    )

    # ── 5. 实验结果 ──
    doc.add_heading('5. 实验结果', level=1)

    doc.add_heading('5.1 分类性能对比', level=2)
    add_styled_table(doc,
        ['排名', '模型', '准确率(%)', 'Macro P(%)', 'Macro R(%)', 'Macro F1(%)', '测试损失'],
        [
            ['1', 'ResNet-50', '99.78', '99.73', '99.81', '99.77', '0.0057'],
            ['2', 'EfficientNet-B3', '99.12', '99.07', '99.31', '99.17', '0.0280'],
            ['3', 'Inception V3', '99.12', '98.89', '99.28', '99.08', '0.0251'],
            ['4', 'VGG-16', '99.12', '99.14', '99.31', '99.21', '0.0267'],
            ['5', 'MobileNet V3', '98.90', '98.87', '99.25', '99.05', '0.0361'],
            ['6', 'VGG-19', '98.68', '98.38', '98.93', '98.62', '0.0560'],
            ['7', 'ViT-B/16', '98.24', '98.21', '98.39', '98.28', '0.0673'],
            ['8', 'Inc-ResNet V2', '94.27', '94.25', '95.36', '94.61', '0.1802'],
            ['9', 'Custom MLP', '63.88', '63.94', '68.69', '64.30', '0.9222'],
        ],
        highlight_best_col=2
    )

    # 性能对比图
    doc.add_paragraph('')
    add_figure(doc, 'model_comparison_bar.png', '图1: 各模型 Accuracy / Macro F1 / Weighted F1 对比')

    doc.add_heading('5.2 关键发现', level=2)
    findings = [
        'ResNet-50以99.78%的测试准确率位居第一，在454张测试图片中仅错分1张。',
        'EfficientNet-B3、Inception V3和VGG-16并列第二，均达到99.12%。',
        '所有预训练CNN模型（排名1-7）均达到98%以上准确率，证明迁移学习的有效性。',
        'Inception-ResNet V2虽然参数量最大（54.70M），准确率仅94.27%，可能在小数据集上过拟合。',
        'Custom MLP（无预训练）仅63.88%，凸显预训练特征提取的重要性。',
    ]
    for f in findings:
        doc.add_paragraph(f, style='List Bullet')

    # 训练曲线
    add_figure(doc, 'training_curves_grid.png', '图2: 所有模型训练曲线网格图（Loss + Accuracy）', width=Inches(6.5))

    # 混淆矩阵
    add_figure(doc, 'confusion_matrices_grid.png', '图3: 所有模型混淆矩阵网格图')

    # 收敛曲线
    add_figure(doc, 'convergence_comparison.png', '图4: 验证集准确率收敛曲线对比')

    add_figure(doc, 'loss_comparison.png', '图5: 验证集损失收敛曲线对比')

    doc.add_heading('5.3 逐类别F1热力图', level=2)
    add_figure(doc, 'per_class_f1_heatmap.png', '图6: 各模型 × 各类别 F1 Score 热力图')

    add_figure(doc, 'per_class_precision_recall.png', '图7: 各模型逐类别 Precision & Recall 分组柱状图', width=Inches(6.5))

    doc.add_heading('5.4 最佳模型详细分析', level=2)
    add_figure(doc, 'best_model_confusion_heatmap.png', '图8: ResNet-50 详细混淆矩阵（计数 + 归一化%）', width=Inches(6.5))

    add_figure(doc, 'best_model_predictions.png', '图9: ResNet-50 推理结果可视化', width=Inches(6.5))

    add_figure(doc, 'error_analysis.png', '图10: 各模型误分类数量分析', width=Inches(6.5))

    # ── 6. 模型效率分析 ──
    doc.add_heading('6. 模型效率分析', level=1)

    doc.add_heading('6.1 计算性能指标', level=2)
    add_styled_table(doc,
        ['模型', '大小(MB)', 'GFLOPs', '延迟(ms)', '吞吐量(FPS)'],
        [
            ['Custom MLP', '49.14', '0.01', '0.31', '3247'],
            ['MobileNet V3', '12.39', '0.23', '7.31', '137'],
            ['EfficientNet-B3', '42.67', '1.93', '14.27', '70'],
            ['ResNet-50', '91.92', '4.13', '6.71', '149'],
            ['VGG-16', '56.65', '15.35', '2.21', '453'],
            ['VGG-19', '76.90', '19.51', '2.71', '369'],
            ['Inception V3', '85.28', '5.75', '14.42', '69'],
            ['ViT-B/16', '328.06', '11.29', '6.07', '165'],
            ['Inc-ResNet V2', '208.93', '13.15', '35.00', '29'],
        ]
    )

    add_figure(doc, 'inference_speed_comparison.png', '图11: 推理延迟与吞吐量对比', width=Inches(6.5))

    add_figure(doc, 'model_size_comparison.png', '图12: 模型大小与参数量对比', width=Inches(6.5))

    add_figure(doc, 'flops_vs_accuracy.png', '图13: GFLOPs vs 准确率（气泡大小 ∝ 模型大小）')

    add_figure(doc, 'params_vs_accuracy.png', '图14: 参数量 vs 准确率散点图')

    add_figure(doc, 'training_time_vs_accuracy.png', '图15: 训练时间 vs 准确率散点图')

    add_figure(doc, 'efficiency_radar.png', '图16: 多维效率雷达图')

    doc.add_heading('6.2 性价比分析', level=2)
    add_styled_table(doc,
        ['模型', '准确率(%)', '大小(MB)', 'GFLOPs', '综合评价'],
        [
            ['ResNet-50', '99.78', '91.92', '4.13', '精度最高，推荐首选'],
            ['MobileNet V3', '98.90', '12.39', '0.23', '最轻量，推荐部署'],
            ['EfficientNet-B3', '99.12', '42.67', '1.93', '最佳性价比'],
            ['VGG-16', '99.12', '56.65', '15.35', '吞吐量最高(453 FPS)'],
        ]
    )

    # ── 7. 消融实验 ──
    doc.add_heading('7. 消融实验', level=1)
    doc.add_paragraph(
        '对主要模型进行模块级消融实验：逐个将模型不同模块的参数置零，'
        '观察准确率下降程度以衡量各模块对分类性能的贡献。'
    )

    add_figure(doc, 'ablation_study.png', '图17: 消融实验 — 各模块贡献分析', width=Inches(6.5))

    add_figure(doc, 'module_contribution_heatmap.png', '图18: 跨模型模块贡献比热力图', width=Inches(6.5))

    doc.add_paragraph(
        '消融实验表明：所有模块消融后准确率均大幅下降至随机水平附近（12%~28%），'
        '说明模型各模块紧密协作。浅层特征提取的消融影响与深层模块同样严重，'
        '表明古建筑文字分类同时依赖低级纹理特征和高级语义特征。'
    )

    # ── 8. Grad-CAM 可视化 ──
    doc.add_heading('8. Grad-CAM 可视化分析', level=1)
    doc.add_paragraph(
        '采用 Grad-CAM 技术对各模型的中间层进行可视化分析，'
        '生成注意力热力图以直观展示模型分类时关注的图像区域。'
    )

    fig_num = 19
    gradcam_models = [
        ('resnet50', 'ResNet-50'),
        ('vgg16', 'VGG-16'),
        ('vgg19', 'VGG-19'),
        ('inception_v3', 'Inception V3'),
        ('inception_resnet_v2', 'Inception-ResNet V2'),
        ('efficientnet_b3', 'EfficientNet-B3'),
        ('mobilenet_v3', 'MobileNet V3'),
        ('vit_b16', 'ViT-B/16'),
    ]

    for model_key, model_name in gradcam_models:
        fname = f'gradcam_layers_{model_key}.png'
        if os.path.exists(os.path.join(FIG_DIR, fname)):
            add_figure(doc, fname, f'图{fig_num}: {model_name} 各层 Grad-CAM 热力图', width=Inches(6.5))
            fig_num += 1

    add_figure(doc, 'gradcam_best_model_grid.png',
               f'图{fig_num}: ResNet-50 多样本 Grad-CAM 可视化', width=Inches(6.5))
    fig_num += 1

    # ── 9. 通道激活分析 ──
    doc.add_heading('9. 通道激活热力图', level=1)
    doc.add_paragraph(
        '对各模型中间层的通道激活进行统计和可视化，展示 Top-8 最强激活通道的空间分布。'
    )

    for model_key, model_name in gradcam_models:
        fname = f'channel_heatmap_{model_key}.png'
        if os.path.exists(os.path.join(FIG_DIR, fname)):
            add_figure(doc, fname, f'图{fig_num}: {model_name} 通道激活热力图', width=Inches(6.5))
            fig_num += 1

    # ── 10. 纹理分析与局部细节对比 ──
    doc.add_heading('10. 纹理分析与局部细节对比', level=1)

    doc.add_heading('10.1 六类建筑纹理特征对比', level=2)
    doc.add_paragraph(
        '对六类古建筑图像分别提取边缘（Sobel算子）、高频响应（Laplacian算子）和纹理复杂度（局部方差）等纹理特征，'
        '横向对比不同类别建筑在纹理结构上的差异。'
    )
    add_figure(doc, 'texture_multi_class_comparison.png',
               f'图{fig_num}: 六类古建筑纹理特征横向对比', width=Inches(6.5))
    fig_num += 1

    doc.add_heading('10.2 逐类别局部放大细节分析', level=2)
    doc.add_paragraph(
        '对每个类别的代表性样本，自动检测高纹理复杂度的感兴趣区域（ROI），'
        '并对原图、边缘检测图、纹理复杂度图分别进行局部放大，展示各区域的细微纹理差异。'
    )

    for cls_idx in range(6):
        fname = f'texture_detail_zoom_class{cls_idx}.png'
        if os.path.exists(os.path.join(FIG_DIR, fname)):
            add_figure(doc, fname,
                       f'图{fig_num}: 类别 {cls_idx + 1} 局部放大纹理细节对比', width=Inches(6.5))
            fig_num += 1

    doc.add_heading('10.3 纹理特征深度分析', level=2)
    doc.add_paragraph(
        '选取不同类别的样本进行深度纹理分析：标注感兴趣区域后，对边缘响应、梯度方向、'
        '高频Laplacian响应、纹理复杂度分别进行ROI放大可视化，并统计ROI区域的特征值分布。'
    )
    add_figure(doc, 'texture_feature_analysis.png',
               f'图{fig_num}: 纹理特征深度分析与局部放大', width=Inches(6.5))
    fig_num += 1

    doc.add_heading('10.4 跨模型注意力局部放大对比', level=2)
    doc.add_paragraph(
        '对同一输入图像，使用6个模型（ResNet-50、VGG-16、Inception V3、EfficientNet-B3、'
        'MobileNet V3、ViT-B/16）生成Grad-CAM热力图，标注高纹理ROI区域后进行局部放大，'
        '直观对比不同模型在局部细节区域的注意力差异。'
    )
    add_figure(doc, 'texture_cross_model_zoom.png',
               f'图{fig_num}: 跨模型注意力局部放大对比', width=Inches(6.5))
    fig_num += 1

    doc.add_heading('10.5 最佳模型 vs 对比模型注意力对比', level=2)
    doc.add_paragraph(
        '选取最佳模型（ResNet-50）与对比模型（ViT-B/16）进行注意力对比分析。'
        '对4类不同建筑样本分别展示全图Grad-CAM热力图和ROI放大视图，'
        '同时叠加纹理边缘图进行综合分析，揭示不同模型在纹理细节区域的关注差异。'
    )
    add_figure(doc, 'texture_attention_zoom.png',
               f'图{fig_num}: 最佳模型 vs 对比模型注意力与纹理局部放大对比', width=Inches(6.5))
    fig_num += 1

    doc.add_paragraph(
        '从跨模型对比可以看出，ResNet-50和VGG-16在建筑纹理细节（如瓦片、梁柱纹路）上的注意力更为集中精准，'
        '而ViT-B/16的注意力分布更加分散，倾向于关注全局结构而非局部纹理。'
        'EfficientNet-B3和MobileNet V3则在两者之间取得了较好的平衡。'
    )

    # ── 11. 结论 ──
    doc.add_heading('11. 实验结论', level=1)

    doc.add_heading('11.1 主要结论', level=2)
    conclusions = [
        '迁移学习在古建筑文字分类中高度有效：所有预训练模型均达到94%以上准确率，大幅优于从头训练的Custom MLP（63.88%）。',
        'ResNet-50是最佳模型选择：99.78%准确率、99.77% Macro F1，计算量适中（4.13 GFLOPs）。',
        '轻量级部署推荐MobileNet V3：仅12.39MB、0.23 GFLOPs，仍达98.90%准确率。',
        '高效平衡之选为EfficientNet-B3：42.67MB、1.93 GFLOPs，准确率99.12%。',
        '更大的模型不一定更好：ViT-B/16（86M参数）和Inception-ResNet V2（54.7M参数）准确率反而较低。',
        '古建筑文字分类依赖多层级特征：消融实验表明浅层和深层特征同等重要。',
    ]
    for i, c in enumerate(conclusions, 1):
        doc.add_paragraph(f'{i}. {c}')

    doc.add_heading('11.2 模型推荐', level=2)
    add_styled_table(doc,
        ['应用场景', '推荐模型', '准确率', '理由'],
        [
            ['追求最高精度', 'ResNet-50', '99.78%', '精度最高，仅1例错误'],
            ['移动端/嵌入式', 'MobileNet V3', '98.90%', '最小模型，最低计算量'],
            ['精度-效率均衡', 'EfficientNet-B3', '99.12%', '最佳性价比'],
            ['实时处理场景', 'VGG-16', '99.12%', '最高吞吐量(453 FPS)'],
        ]
    )

    doc.add_heading('11.3 局限性与未来工作', level=2)
    limitations = [
        '数据集规模有限（2,269张），扩大数据规模可进一步提升泛化能力',
        '类别不平衡问题可考虑更多数据增强或过采样策略',
        'Inception-ResNet V2可通过调整超参数优化性能',
        '未来可尝试Top-3模型集成以进一步提升性能',
    ]
    for lim in limitations:
        doc.add_paragraph(lim, style='List Bullet')

    # 综合表格
    doc.add_page_break()
    doc.add_heading('附录: 综合性能对比表', level=1)
    add_figure(doc, 'comprehensive_performance_table.png', '附表: 含效率指标的完整对比表', width=Inches(6.5))

    # ── 保存 ──
    out_path = os.path.join(PAPER_DIR, '实验总结报告.docx')
    doc.save(out_path)
    print(f'Saved: {out_path}')


if __name__ == '__main__':
    main()
