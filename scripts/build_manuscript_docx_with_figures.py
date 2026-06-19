#!/usr/bin/env python
"""Build the English AAFNet manuscript DOCX with curated experiment figures.

The source manuscript intentionally keeps figures in an appendix index instead
of inline Markdown image tags. This script creates a reproducible assembled
Markdown copy with key figure panels inserted near their related sections, then
calls pandoc to produce a Word document.
"""

from __future__ import annotations

import argparse
import re
import shutil
import subprocess
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_IN = ROOT / "paper" / "manuscript.md"
DEFAULT_MD = ROOT / "paper" / "manuscript_with_figures.md"
DEFAULT_DOCX = ROOT / "paper" / "AAFNet_manuscript_with_figures.docx"
DEFAULT_AUDIT = ROOT / "paper" / "figure_formula_audit.md"


FIGURES: dict[str, list[tuple[str, str, str]]] = {
    "# AAFNet: Architectural-Aware Multi-Scale Fusion for Robust Ancient-Building Image Classification": [
        (
            "figures/F_O_hero.png",
            "Figure F-O. Six-panel summary of AAFNet's main evidence chain.",
            "95%",
        ),
    ],
    "### 3.1 AAFNet architecture overview": [
        (
            "figures/diagrams/D1_aafnet_arch.png",
            "Figure D1. AAFNet end-to-end architecture.",
            "95%",
        ),
    ],
    "### 4.2 Data Audit": [
        (
            "figures/F_TC0_three_corpus_samples.png",
            "Figure F-TC0. Representative image samples from the three analysis corpora.",
            "95%",
        ),
        (
            "figures/F_TC1_three_corpus_audit.png",
            "Figure F-TC1. Three-corpus data audit, split size, class difficulty, and deduplication retention.",
            "95%",
        ),
        (
            "figures/F_P_strict_eval.png",
            "Figure F-P. Strict-test re-evaluation after audit.",
            "90%",
        ),
    ],
    "### 5.1 Main comparison (T2)": [
        (
            "figures/F_N_pareto_efficiency.png",
            "Figure F-N. Accuracy-versus-compute Pareto comparison.",
            "90%",
        ),
    ],
    "### 5.2 Robustness — component attribution (Figure 4 + T7)": [
        (
            "figures/F_A_attrib_heatmap.png",
            "Figure F-A. Seven-cell robustness attribution heatmap.",
            "95%",
        ),
        (
            "figures/F_C_robustness_grouped_bar.png",
            "Figure F-C. Robustness grouped bar across perturbation families.",
            "95%",
        ),
    ],
    "### 5.2.4 Rotation robustness — 24-angle closed-loop sweep (T7.2)": [
        (
            "figures/F_AA_rotation_polar.png",
            "Figure F-AA. Closed-loop polar accuracy over 24 rotation angles.",
            "80%",
        ),
    ],
    "### 5.3.2 AAFNet on the harder corpora — ASP_clean and AS25_clean": [
        (
            "figures/F_M_cross_corpus.png",
            "Figure F-M. Cross-corpus baseline-vs-AAFNet accuracy.",
            "85%",
        ),
        (
            "figures/F_TC2_three_corpus_performance.png",
            "Figure F-TC2. Three-corpus performance, macro-F1, effect direction, and strict-test stability.",
            "95%",
        ),
        (
            "figures/F_TC3_three_corpus_confusions.png",
            "Figure F-TC3. Row-normalized confusion matrices for AL6, ASP_clean, and AS25_clean.",
            "95%",
        ),
        (
            "figures/F_TC4_three_corpus_training.png",
            "Figure F-TC4. Reference training dynamics for the three analysis corpora.",
            "95%",
        ),
        (
            "figures/F_TC6_three_corpus_parity.png",
            "Figure F-TC6. Seed-42 three-corpus parity for robustness, calibration, and rotation probes.",
            "95%",
        ),
    ],
    "### 5.5 Additional baselines and multi-model significance": [
        (
            "figures/F_L_friedman_cv.png",
            "Figure F-L. Friedman CV bar and p-value matrix.",
            "95%",
        ),
    ],
    "### 5.6 Calibration with temperature scaling": [
        (
            "figures/F_E_calibration_dual.png",
            "Figure F-E. Calibration ECE before and after temperature scaling.",
            "95%",
        ),
    ],
    "### 5.7 Feature-quality probes": [
        (
            "figures/F_TC5_three_corpus_probe_matrix.png",
            "Figure F-TC5. Three-corpus matrix of completed probe deltas.",
            "95%",
        ),
        (
            "figures/F_Q_downstream_radar.png",
            "Figure F-Q. Downstream-battery radar summary.",
            "85%",
        ),
        (
            "figures/downstream_inference/F_Q2a_retrieval_effect.png",
            "Figure F-Q2a. AL6 retrieval inference with a query image and model-nearest neighbours.",
            "95%",
        ),
        (
            "figures/downstream_inference/F_Q2b_ood_asp_effect.png",
            "Figure F-Q2b. OOD inference for ASP_clean using AL6-trained energy scores.",
            "95%",
        ),
        (
            "figures/downstream_inference/F_Q2c_ood_as25_effect.png",
            "Figure F-Q2c. OOD inference for AS25_clean using AL6-trained energy scores.",
            "95%",
        ),
        (
            "figures/downstream_inference/F_Q2d_selective_rejection_effect.png",
            "Figure F-Q2d. Selective prediction under Gaussian noise.",
            "95%",
        ),
        (
            "figures/downstream_inference/F_Q2e_augmentation_invariance_effect.png",
            "Figure F-Q2e. Augmentation-invariance inference in feature space.",
            "95%",
        ),
        (
            "figures/downstream_inference/F_Q2f_domain_probe_effect.png",
            "Figure F-Q2f. Domain-probe inference from frozen AL6-trained features.",
            "95%",
        ),
        (
            "figures/downstream_inference/F_Q2g_fewshot_asp_effect.png",
            "Figure F-Q2g. Few-shot ASP_clean inference using frozen-feature prototypes.",
            "95%",
        ),
        (
            "figures/downstream_inference/F_Q2h_fewshot_as25_effect.png",
            "Figure F-Q2h. Few-shot AS25_clean inference using frozen-feature prototypes.",
            "95%",
        ),
    ],
    "### 5.8 Interpretability": [
        (
            "figures/F_TC7_public_interpretability.png",
            "Figure F-TC7. Public-corpus inference traces with input images and predicted-class GradCAM overlays.",
            "95%",
        ),
        (
            "figures/F_R_gradcam_comparison.png",
            "Figure F-R. GradCAM activation comparison.",
            "95%",
        ),
    ],
}


MATH_CODE_BLOCKS: dict[str, str] = {}


def protect_math(text: str) -> tuple[str, list[str]]:
    blocks: list[str] = []

    def repl(match: re.Match[str]) -> str:
        blocks.append(match.group(0))
        return f"@@MATH{len(blocks) - 1}@@"

    protected = re.sub(r"\$\$.*?\$\$|\$[^$\n]+\$", repl, text, flags=re.S)
    return protected, blocks


def restore_math(text: str, blocks: list[str]) -> str:
    for idx, block in enumerate(blocks):
        text = text.replace(f"@@MATH{idx}@@", block)
    return text


def protect_code_spans(text: str) -> tuple[str, list[str]]:
    spans: list[str] = []

    def repl(match: re.Match[str]) -> str:
        spans.append(match.group(0))
        return f"@@CODE{len(spans) - 1}@@"

    return re.sub(r"`[^`\n]+`", repl, text), spans


def restore_code_spans(text: str, spans: list[str]) -> str:
    for idx, span in enumerate(spans):
        text = text.replace(f"@@CODE{idx}@@", span)
    return text


def sub_math(pattern: str, replacement: str, text: str) -> str:
    """Regex replace while preserving TeX backslashes in replacement text."""

    def repl(match: re.Match[str]) -> str:
        out = replacement
        for idx, group in enumerate(match.groups(), start=1):
            out = out.replace(f"\\{idx}", group or "")
        return out

    return re.sub(pattern, repl, text)


def mathify_line(line: str) -> str:
    if not line.strip() or line.startswith("![") or line.startswith("```"):
        return line

    line, math_blocks = protect_math(line)
    line, code_spans = protect_code_spans(line)

    replacements = [
        (
            r"\*\*x\*\* ∈ ℝ\^\{3×H×W\}",
            r"$\mathbf{x} \in \mathbb{R}^{3\times H\times W}$",
        ),
        (
            r"θ ∈ ℝ\^P \(P = ([0-9,]+) at h=([0-9]+), L=([0-9]+)\)",
            r"$\theta \in \mathbb{R}^{P}$ ($P = \1$, $h=\2$, $L=\3$)",
        ),
        (
            r"θ ∈ ℝ\^P",
            r"$\theta \in \mathbb{R}^{P}$",
        ),
        (
            r"W ∈ ℝ\^\{Ci × Co × Do × Di\}",
            r"$W \in \mathbb{R}^{C_i \times C_o \times D_o \times D_i}$",
        ),
        (
            r"χ²\s*\(([0-9]+)\)\s*=\s*([0-9.]+),\s*p\s*=\s*([0-9.]+)",
            r"$\chi^2(\1) = \2,\ p = \3$",
        ),
        (
            r"χ²\s*=\s*([0-9.]+),\s*p\s*=\s*([0-9.]+)",
            r"$\chi^2 = \1,\ p = \2$",
        ),
        (
            r"χ²\s*\(([0-9]+)\)\s*=\s*\*\*([0-9.]+),\s*p < ([0-9.]+)\*\*",
            r"$\chi^2(\1) = \2,\ p < \3$",
        ),
        (
            r"χ² rises from ([0-9.]+) \(5 models\) → ([0-9.]+) \(6",
            r"$\chi^2$ rises from $\1$ (5 models) $\rightarrow$ $\2$ (6",
        ),
        (
            r"σ ∈ \{0\.05, 0\.10, 0\.20\}",
            r"$\sigma \in \{0.05, 0.10, 0.20\}$",
        ),
        (
            r"Δ ∈ \{±0\.2, ±0\.3, ±0\.4\}",
            r"$\Delta \in \{\pm 0.2, \pm 0.3, \pm 0.4\}$",
        ),
        (
            r"θ ∈ \{0°, 15°, 30°, …, 345°\}",
            r"$\theta \in \{0^\circ, 15^\circ, 30^\circ, \ldots, 345^\circ\}$",
        ),
        (
            r"θ ∈ \{0°, 15°, …, 345°\}",
            r"$\theta \in \{0^\circ, 15^\circ, \ldots, 345^\circ\}$",
        ),
        (
            r"5×10⁻⁴",
            r"$5\times 10^{-4}$",
        ),
        (
            r"8 × 10⁻⁴",
            r"$8\times 10^{-4}$",
        ),
        (
            r"PSNR ≥ ([0-9]+) dB",
            r"$\mathrm{PSNR} \ge \1\,\mathrm{dB}$",
        ),
        (
            r"pHash≤6",
            r"pHash $\le 6$",
        ),
        (
            r"pHash ≤ 6",
            r"pHash $\le 6$",
        ),
        (
            r"threshold ≤ 6",
            r"threshold $\le 6$",
        ),
        (
            r"mean_\{j≠k\}",
            r"$\mathrm{mean}_{j\ne k}$",
        ),
    ]
    for pattern, replacement in replacements:
        line = sub_math(pattern, replacement, line)

    line = sub_math(r"σ\s*=\s*([0-9.]+)", r"$\sigma = \1$", line)
    line = sub_math(r"θ\s*=\s*([0-9]+)°", r"$\theta = \1^\circ$", line)
    line = sub_math(r"τ\s*=\s*([0-9.]+)", r"$\tau = \1$", line)
    line = sub_math(r"γ\s*=\s*([0-9.]+),\s*ε\s*=\s*([0-9.]+)", r"$\gamma = \1,\ \epsilon = \2$", line)
    line = sub_math(r"\bT\s*=\s*([0-9]+)\b", r"$T = \1$", line)
    line = sub_math(r"\bα\s*=\s*([0-9.]+)\b", r"$\alpha = \1$", line)
    line = sub_math(r"\bn\s*=\s*([0-9]+)\b", r"$n = \1$", line)
    line = sub_math(r"\bP\s*=\s*([0-9,]+)\b", r"$P = \1$", line)
    line = sub_math(r"\bh\s*=\s*([0-9]+)\b", r"$h = \1$", line)
    line = sub_math(r"\bL\s*=\s*([0-9]+)\b", r"$L = \1$", line)
    line = sub_math(r"\bECE\s*=\s*([0-9.]+)", r"$\mathrm{ECE} = \1$", line)
    line = sub_math(r"\bAUROC\s*=\s*([0-9.]+)", r"$\mathrm{AUROC} = \1$", line)
    line = sub_math(r"\bAURC\s+at\s+([0-9.]+)", r"AURC at $\1$", line)
    line = sub_math(r"(?<![A-Za-z])([+-]?\d+(?:\.\d+)?)\s*±\s*([0-9.]+)\s*%", r"$\1 \pm \2\,\%$", line)
    line = sub_math(r"(\d+)-fold × (\d+)-seed", r"$\1\times\2$-seed", line)
    line = sub_math(r"(\d+) seeds × (\d+) epochs", r"$\1$ seeds $\times$ $\2$ epochs", line)
    line = sub_math(r"(\d+) seeds × (\d+) fold ×", r"$\1$ seeds $\times$ $\2$ fold $\times$", line)
    line = sub_math(r"(\d+) perturbations × (\d+) severities", r"$\1$ perturbations $\times$ $\2$ severities", line)
    line = sub_math(r"(\d+)×(\d+)", r"$\1\times\2$", line)
    line = sub_math(r"(\d+) × (\d+)", r"$\1\times\2$", line)
    line = sub_math(r"Δ\s*=\s*([+−-]?[0-9.]+)\s*pp", r"$\Delta = \1\,\mathrm{pp}$", line)
    line = sub_math(r"Δ\s+AAFNet\s*−\s*baseline", r"$\Delta$ AAFNet $-$ baseline", line)
    line = sub_math(r"Δ\s+vs\s+reference", r"$\Delta$ vs reference", line)
    line = sub_math(r"\| Δ \|", r"| $\Delta$ |", line)
    line = sub_math(r"\| Δ$", r"| $\Delta$", line)

    line = restore_code_spans(line, code_spans)
    line = restore_math(line, math_blocks)
    return line


def mathify_text(text: str) -> str:
    for old, new in MATH_CODE_BLOCKS.items():
        text = text.replace(old, new)

    out: list[str] = []
    in_code = False
    for line in text.splitlines():
        if line.startswith("```"):
            in_code = not in_code
            out.append(line)
            continue
        if in_code:
            out.append(line)
        else:
            out.append(mathify_line(line))
    return normalize_display_math("\n".join(out) + "\n")


def normalize_display_math(text: str) -> str:
    """Use \[...\] for display math to avoid literal $$ in DOCX output."""
    out: list[str] = []
    in_display = False
    for line in text.splitlines():
        if line.strip() == "$$":
            out.append(r"\[" if not in_display else r"\]")
            in_display = not in_display
        else:
            out.append(line)
    return "\n".join(out) + "\n"


def figure_block(figures: list[tuple[str, str, str]], paper_dir: Path) -> str:
    lines = ["", "<!-- auto-inserted figure block -->", ""]
    for rel_path, caption, width in figures:
        path = paper_dir / rel_path
        if not path.exists():
            lines.append(f"> **Missing figure:** `{rel_path}`")
            lines.append("")
            continue
        lines.append(f"![{caption}]({rel_path}){{width={width}}}")
        lines.append("")
    return "\n".join(lines)


def build_markdown(input_path: Path, output_path: Path) -> int:
    paper_dir = output_path.parent
    text = mathify_text(input_path.read_text(encoding="utf-8"))
    lines = text.splitlines()
    out: list[str] = []
    inserted = 0
    for line in lines:
        out.append(line)
        figs = FIGURES.get(line.strip())
        if figs:
            out.append(figure_block(figs, paper_dir))
            inserted += len(figs)
    output_path.write_text("\n".join(out) + "\n", encoding="utf-8")
    return inserted


def audit_figures(md_path: Path, docx_path: Path, audit_path: Path) -> None:
    from PIL import Image
    from docx import Document

    md = md_path.read_text(encoding="utf-8")
    figure_refs = re.findall(r"!\[([^\]]*)\]\(([^)]+)\)\{width=([^}]+)\}", md)
    rows: list[str] = []
    ok_count = 0
    missing_count = 0
    for idx, (caption, rel_path, width) in enumerate(figure_refs, start=1):
        path = md_path.parent / rel_path
        if path.exists():
            try:
                with Image.open(path) as im:
                    size = f"{im.size[0]}x{im.size[1]}"
                status = "OK"
                ok_count += 1
            except Exception as exc:  # pragma: no cover - diagnostic path
                size = f"unreadable: {exc}"
                status = "BAD"
        else:
            size = "-"
            status = "MISSING"
            missing_count += 1
        rows.append(
            f"| {idx} | {status} | `{rel_path}` | {width} | {size} | {caption} |"
        )

    inline_shapes = 0
    tables = 0
    paragraphs = 0
    if docx_path.exists():
        doc = Document(docx_path)
        inline_shapes = len(doc.inline_shapes)
        tables = len(doc.tables)
        paragraphs = len(doc.paragraphs)

    report = [
        "# Figure and Formula Audit",
        "",
        f"- Assembled Markdown: `{md_path.relative_to(ROOT)}`",
        f"- DOCX: `{docx_path.relative_to(ROOT)}`",
        f"- Referenced figures: {len(figure_refs)}",
        f"- Readable figures: {ok_count}",
        f"- Missing figures: {missing_count}",
        f"- DOCX inline shapes: {inline_shapes}",
        f"- DOCX tables: {tables}",
        f"- DOCX paragraphs: {paragraphs}",
        "",
        "## Figure Checklist",
        "",
        "| # | Status | File | Width | Pixel size | Caption |",
        "|---:|---|---|---:|---:|---|",
        *rows,
        "",
        "## Formula Rendering",
        "",
        "Formula rendering is verified after DOCX generation by checking for OMML",
        "math nodes in `word/document.xml`; see the command output recorded in",
        "the build/QA step.",
        "",
    ]
    audit_path.write_text("\n".join(report), encoding="utf-8")


def run_pandoc(input_md: Path, output_docx: Path) -> None:
    pandoc = shutil.which("pandoc")
    if not pandoc:
        raise RuntimeError("pandoc is not available on PATH")

    cmd = [
        pandoc,
        str(input_md),
        "--from=markdown+tex_math_dollars+tex_math_single_backslash+pipe_tables+raw_html+link_attributes",
        "--to=docx",
        "--standalone",
        "--toc",
        f"--resource-path={ROOT}:{input_md.parent}",
        "--metadata",
        "title=AAFNet: An Architectural-Aware Fusion Network",
        "-o",
        str(output_docx),
    ]
    subprocess.run(cmd, cwd=ROOT, check=True)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--in", dest="input", default=str(DEFAULT_IN))
    parser.add_argument("--md-out", default=str(DEFAULT_MD))
    parser.add_argument("--docx-out", default=str(DEFAULT_DOCX))
    parser.add_argument("--audit-out", default=str(DEFAULT_AUDIT))
    parser.add_argument("--skip-pandoc", action="store_true")
    args = parser.parse_args()

    input_path = Path(args.input).resolve()
    md_out = Path(args.md_out).resolve()
    docx_out = Path(args.docx_out).resolve()
    audit_out = Path(args.audit_out).resolve()

    md_out.parent.mkdir(parents=True, exist_ok=True)
    inserted = build_markdown(input_path, md_out)

    if not args.skip_pandoc:
        run_pandoc(md_out, docx_out)
    audit_figures(md_out, docx_out, audit_out)

    print(f"Wrote assembled Markdown: {md_out}")
    if not args.skip_pandoc:
        print(f"Wrote DOCX: {docx_out}")
    print(f"Wrote audit: {audit_out}")
    print(f"Inserted figures: {inserted}")


if __name__ == "__main__":
    main()
